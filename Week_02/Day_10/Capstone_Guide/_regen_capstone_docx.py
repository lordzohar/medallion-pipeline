"""Regenerate Capstone.docx with the actual real-streams capstone content.

Preserves the original Word-document structure (Title, Heading 1 sections,
List Bullet / List Number) so the document still looks consistent with the
course style. Overwrites Capstone.docx in place.
"""
from pathlib import Path
import docx

HERE = Path(__file__).resolve().parent
OUT  = HERE / "Capstone.docx"

doc = docx.Document()

def H(text, level=1): doc.add_heading(text, level=level)
def P(text):          doc.add_paragraph(text)
def CODE(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    for r in p.runs: r.font.name = "Consolas"
def UL(items):
    for it in items: doc.add_paragraph(it, style="List Bullet")
def OL(items):
    for it in items: doc.add_paragraph(it, style="List Number")

# ---- Title ---------------------------------------------------------------
t = doc.add_paragraph()
t.style = doc.styles["Title"]
t.add_run("Day 10 Student Lab Guide")
P("Capstone: Real-Time Medallion Pipeline with CDC, Kafka, MinIO, Apache Hop, Airflow and Live Dashboards")
P("")

# ---- Capstone scenario ---------------------------------------------------
H("Capstone Scenario")
P(
    "Build and operate an end-to-end real-time data pipeline that ingests THREE genuine public "
    "streaming sources, lands them in object storage as a medallion (bronze / silver / gold), "
    "exposes the gold layer through dashboards, and is fully observable. The pipeline must include "
    "Postgres CDC via Debezium so reference / configuration changes flow through the same medallion "
    "as the streams. A glidertracker.org-style live map is the user-facing 'wow' surface."
)
P("Sources you will integrate:")
UL([
    "Open Glider Network (OGN) — APRS-IS TCP feed of live aircraft positions.",
    "NOAA api.weather.gov — REST polling of US weather station observations + alerts.",
    "EMSC Seismic Portal — WebSocket of real earthquake events worldwide.",
    "Postgres reference tables (regions, alert_thresholds, subscriber_watchlist) via Debezium pgoutput.",
])

# ---- Architecture roles --------------------------------------------------
H("Where Debezium, Airflow, and the Live Dashboard Fit")
P(
    "The public feeds are the real-time data sources, but the capstone also "
    "needs control-plane data, orchestration, and user-facing views. These "
    "parts are separate on purpose:"
)
UL([
    "Debezium captures changes from the Postgres config database. It is not a fourth sensor feed; it streams reference/config changes such as regions, alert thresholds, and subscriber watchlists into Kafka as CDC topics.",
    "Airflow orchestrates batch-style work around the live streams. The ingestors run continuously as services, while Airflow schedules config drift, bronze-to-silver promotion, silver-to-gold marts, data quality checks, and business dashboard refreshes.",
    "The live-map dashboard consumes the OGN, NOAA, and EMSC Kafka topics directly so the demo can show movement immediately. It does not wait for the 5-minute silver/gold batch cycle.",
    "The business dashboard reads gold Parquet marts. The quality dashboard reads rule results, DLQ state, and Alertmanager webhooks.",
])
P("One-screen flow:")
CODE(
    "OGN / NOAA / EMSC -> ingestor services -> Kafka -> live-map-dashboard\n"
    "config-db Postgres -> Debezium -> Kafka config.public.* topics\n"
    "Kafka -> S3 Sink connectors -> MinIO bronze Avro\n"
    "Airflow -> Hop/Python transforms -> MinIO silver Avro -> MinIO gold Parquet\n"
    "gold Parquet -> business-dashboard; quality results / alerts -> quality-dashboard\n"
    "Prometheus scrapes Kafka, Postgres, MinIO, Connect, Airflow-facing services, and dashboards\n"
)

# ---- Deliverables --------------------------------------------------------
H("Deliverables")
UL([
    "A running docker-compose stack (~24 containers) on a single bridge network.",
    "A scoped demo city (default: Denver / Boulder, CO) configured through a single .env block.",
    "Avro contracts in Schema Registry for the 4 stream subjects.",
    "Debezium source + two S3 Sink connectors (CDC topics, stream topics) configured and RUNNING.",
    "Bronze landing in MinIO under s3://bronze/<topic>/year=…/month=…/day=…/hour=…/ as Avro.",
    "Silver layer (Avro) produced by Apache Hop pipelines or their Python reference transforms.",
    "Gold layer (Parquet) with FOUR marts: aircraft_density_by_region, weather_snapshot, "
    "seismic_24h_summary, region_alert_correlation.",
    "Airflow DAGs orchestrating the medallion every 5 min plus a config-drift DAG.",
    "Three Flask dashboards: quality (DataOps), business (gold parquet), live map (Leaflet + Socket.IO).",
    "Prometheus + Alertmanager + Grafana with at least one freshness rule + one DLQ rule.",
    "A passing smoke test (tests/smoke_test.py) and a written demo script.",
])

# ---- Task 1 --------------------------------------------------------------
H("Task 1 — Start the Stack")
P("All commands are run from the capstone folder:")
CODE("cd Week_02\\Day_10\\Capstone_Files\nCopy-Item .env.example .env\n.\\bootstrap.ps1                  # macOS/Linux: bash bootstrap.sh\n")
P(
    "bootstrap.ps1 is idempotent — re-run it any time. It downloads Kafka Connect plugins + the "
    "JMX agent, brings the stack up, waits for MinIO / Schema Registry / Kafka Connect / Airflow to "
    "become healthy, creates the three MinIO buckets, registers the 4 Avro schemas, registers the "
    "Debezium source + two S3 Sink connectors, unpauses the 4 continuous DAGs, then prints the URL list."
)
P("Verify everything is up:")
CODE("docker compose ps\npython tests\\smoke_test.py\n")

# ---- Task 2 --------------------------------------------------------------
H("Task 2 — Configure the Demo City")
P(
    "Every source is scoped to one city to keep volumes small. Open .env and edit the DEMO CITY block. "
    "Default is Denver / Boulder, CO (real glider activity at Boulder + 6 NOAA stations within 250 km)."
)
CODE(
    "CITY_NAME=Denver\n"
    "CITY_LAT=39.74\n"
    "CITY_LON=-104.99\n"
    "CITY_RADIUS_KM=250\n"
    "NOAA_STATIONS=KDEN,KBJC,KAPA,KCOS,KFNL,KGJT\n"
)
P("Presets pre-written in .env.example comments: New York (KJFK,KLGA,KEWR…), Reno/Minden NV (KRNO,KMEV…).")
P("If you change the city after the stack is already up, recreate only the affected services:")
CODE("docker compose --env-file .env up -d --force-recreate ingestor-ogn ingestor-noaa live-map-dashboard\n")

# ---- Task 3 --------------------------------------------------------------
H("Task 3 — Verify Stream Ingestion (Bronze)")
P("Three long-running ingestor services produce Avro into Kafka:")
UL([
    "ingestor-ogn  → topic ogn.aircraft.positions",
    "ingestor-noaa → topics noaa.observations and noaa.alerts",
    "ingestor-seismic → topic seismic.events",
])
P("Watch them populate in Kafka UI (http://localhost:8088). Then confirm landing in MinIO:")
CODE(
    "# Kafka REST view of connector status\n"
    "curl -s http://localhost:8083/connectors | jq\n"
    "curl -s http://localhost:8083/connectors/pipeline-config-source/status | jq\n"
    "curl -s http://localhost:8083/connectors/s3-sink-bronze-streams/status | jq\n"
)
P("Browse MinIO console (http://localhost:9001, minioadmin / minioadmin) → bucket 'bronze' → see hourly-partitioned Avro files.")
P("Capture evidence: one screenshot of Kafka UI topic with messages, one of MinIO bronze tree, one of /connectors REST output.")

# ---- Task 4 --------------------------------------------------------------
H("Task 4 — Run the Medallion (Silver and Gold)")
P("Airflow DAG 30_hop_medallion runs every 5 minutes; you can also trigger it manually from the Airflow UI (http://localhost:8080, airflow / airflow).")
P("It fans out:")
UL([
    "bronze → silver for 4 stream entities + 3 CDC entities (Apache Hop pipeline bronze_to_silver_*.hpl, "
    "or the Python reference transform in hop/transforms/bronze_to_silver.py).",
    "silver → gold for all 4 marts (silver_to_gold_mart.hpl with a MART parameter, or silver_to_gold.py).",
])
P("Inspect the result in MinIO:")
CODE(
    "# silver — Avro, deduped, late-filtered\n"
    "s3://silver/<entity>/year=YYYY/month=MM/day=DD/part-<hash>.avro\n"
    "# gold — Parquet, one immutable snapshot + one overwritten 'latest'\n"
    "s3://gold/<mart>/snapshot=YYYYMMDDTHHMMSSZ/part-0.parquet\n"
    "s3://gold/<mart>/latest.parquet\n"
)
P("Force-rebuild any single mart from the host:")
CODE("docker exec hop python3 /files/project/transforms/silver_to_gold.py --mart=seismic_24h_summary\n")

# ---- Task 5 --------------------------------------------------------------
H("Task 5 — Stakeholder View (Live Map + Business Dashboard)")
P("Two surfaces target two audiences.")
P("(a) Live map — http://localhost:5003 — built directly from the three Kafka stream topics. Open it and walk through:")
OL([
    "Aircraft appear as rotated plane silhouettes with a yellow callsign tag (glidertracker.org style). Click one — it turns magenta and its trail polyline draws live.",
    "Weather tab: NOAA chips (e.g. KDEN 22°C) populate after ~120 s. Click for the full popup (temp / wind / humidity / pressure / visibility).",
    "Seismic tab: events render as filled circles sized + coloured by magnitude. Anything inside the city radius gets a magenta outline + permanent pulse + 'nearby' counter.",
    "Toggle 'Auto-zoom new quakes' so the demo audience sees a quake fly in from anywhere on Earth.",
    "Switch basemap to Esri satellite for the reveal.",
])
P("(b) Business dashboard — http://localhost:5002 — reads gold parquet. Shows the four marts side-by-side. Each mart cell is your Gold-level KPI artefact.")
P("Pick one KPI for your write-up — recommended choices:")
UL([
    "aircraft_density_by_region.positions_last_1h — 'gliders tracked in the last hour, per region'.",
    "weather_snapshot.temperature_c — 'most recent NOAA observation per station'.",
    "seismic_24h_summary.events with magnitude_bucket — 'earthquakes in the last 24h by magnitude band'.",
    "region_alert_correlation.subscribers_notified — 'subscribers whose region just received a quake above the live threshold'.",
])
P("Document the KPI in ONE sentence. Explain freshness: which Kafka event starts the update, and where the user can see the timestamp.")

# ---- Task 6 --------------------------------------------------------------
H("Task 6 — Show the CDC Loop (Debezium → Silver → Gold)")
P("Trigger a live config change and watch it propagate through every layer.")
CODE(
    "# Lower the seismic alert threshold so more quakes will trigger correlation.\n"
    "docker exec config-db psql -U postgres -d config -c \\\n"
    "  \"UPDATE alert_thresholds SET threshold=5.0 WHERE source='seismic' AND metric='magnitude' AND op='>=';\"\n"
)
P(
    "Within ~5 min the change flows Postgres WAL → Debezium → config.public.alert_thresholds → "
    "bronze (Avro) → silver/alert_thresholds/ → region_alert_correlation mart. Refresh the business "
    "dashboard and screenshot the 'before' and 'after' rows."
)
P("The 15_config_drift DAG runs every 2 min and produces small random mutations on the same tables so the CDC stream is always warm during a demo — pause it from Airflow UI if you want quiet topics.")

# ---- Task 7 --------------------------------------------------------------
H("Task 7 — Operational Readiness")
P("Inject one controlled failure each and capture evidence in the quality dashboard (http://localhost:5001) and Alertmanager (http://localhost:9093).")
UL([
    "Stop an ingestor: docker stop ingestor-seismic. After 5–10 min the seismic.freshness rule turns FAIL and the IngestorStale alert fires.",
    "Cause a DLQ event: docker exec config-db psql -U postgres -d config -c \"ALTER TABLE regions ADD COLUMN broken_col TEXT;\" — watch the DLQ panel grow, then revert with DROP COLUMN.",
    "Pause a DAG in Airflow → notice silver and gold for that entity stop refreshing on the quality dashboard timestamps.",
])
P("State for each failure: the monitoring signal, the recovery action, and how long until normal service resumed (visible in the dashboard).")

# ---- Architecture appendix -----------------------------------------------
H("Appendix A — Architecture Reference")
P(
    "This section replaces the standalone ARCHITECTURE.md. It is the compact "
    "source of truth for the implemented stack."
)
H("Design Goals", level=2)
UL([
    "Real streams, not simulators: OGN, NOAA, and EMSC exercise TCP, REST polling, and WebSocket ingestion patterns.",
    "CDC where CDC belongs: Postgres holds config/reference data only; Debezium streams row changes into Kafka.",
    "Medallion in object storage: bronze and silver are Avro; gold is Parquet for fast dashboard reads.",
    "Hop is the canonical transform design; Python reference transforms make the same logic runnable from Airflow and tests.",
    "Three dashboards serve three audiences: live map for source-level proof, quality dashboard for DataOps, business dashboard for gold marts.",
    "Everything observable: Prometheus scrapes Kafka, Postgres, MinIO, Connect, and dashboards; Alertmanager posts alerts to the quality dashboard.",
])
H("Service Inventory", level=2)
UL([
    "Kafka layer: zookeeper, kafka, schema-registry, connect, kafka-ui.",
    "Source/config layer: config-db, postgres-exporter, app-base, app, ingestor-ogn, ingestor-noaa, ingestor-seismic.",
    "Storage/transform layer: minio, hop.",
    "Orchestration layer: airflow-db, airflow-init, airflow-webserver, airflow-scheduler.",
    "Presentation layer: quality-dashboard, business-dashboard, live-map-dashboard.",
    "Observability layer: prometheus, alertmanager, grafana.",
])
H("Topic and Subject Inventory", level=2)
UL([
    "ogn.aircraft.positions — OGN aircraft positions, Avro subject ogn.aircraft.positions-value.",
    "noaa.observations — NOAA station observations, Avro subject noaa.observations-value.",
    "noaa.alerts — NOAA weather alerts, Avro subject noaa.alerts-value.",
    "seismic.events — EMSC earthquake events, Avro subject seismic.events-value.",
    "config.public.regions, config.public.alert_thresholds, config.public.subscriber_watchlist — Debezium-managed CDC topics.",
    "config.heartbeat and dlq.* topics — connector heartbeat and dead-letter queues.",
])
H("Medallion Contracts", level=2)
CODE(
    "Bronze Avro: s3://bronze/<topic>/year=YYYY/month=MM/day=DD/hour=HH/<topic>+<part>+<offset>.avro\n"
    "Silver Avro: s3://silver/<entity>/year=YYYY/month=MM/day=DD/part-<hash>.avro\n"
    "Silver audit: s3://silver/_quality/<entity>/<utc-iso>.avro\n"
    "Gold snapshot: s3://gold/<mart>/snapshot=YYYYMMDDTHHMMSSZ/part-0.parquet\n"
    "Gold latest: s3://gold/<mart>/latest.parquet\n"
)
P("Gold marts:")
UL([
    "aircraft_density_by_region — last-1h aircraft counts and movement stats by region.",
    "weather_snapshot — latest NOAA observation per station joined to region metadata.",
    "seismic_24h_summary — earthquake counts bucketed by magnitude and region.",
    "region_alert_correlation — quakes over the live threshold joined to subscribed regions and watchlists.",
])
H("DAG Topology", level=2)
CODE(
    "00_bootstrap        manual     create topics, buckets, schemas, connectors\n"
    "15_config_drift     every 2m   mutate config tables to keep Debezium CDC warm\n"
    "30_hop_medallion    every 5m   bronze -> silver, then silver -> 4 gold marts\n"
    "40_data_quality     every 5m   rule pack -> quality-dashboard\n"
    "50_business_kpis    every 5m   POST business-dashboard /api/refresh\n"
)
P(
    "The stream ingestors and live-map dashboard are long-running services, "
    "not DAGs. They stay connected to Kafka continuously."
)
H("Observability Map", level=2)
CODE(
    "Kafka JMX :9404, postgres-exporter :9187, MinIO :9000,\n"
    "quality-dashboard :5001, business-dashboard :5002,\n"
    "live-map-dashboard :5003, and kafka-connect REST :8083\n"
    "        -> Prometheus :9090 -> Alertmanager :9093\n"
    "        -> quality-dashboard /webhook/alerts\n"
)

# ---- Runbook appendix -----------------------------------------------------
H("Appendix B — Runbook Quick Reference")
P(
    "This section replaces the standalone RUNBOOK.md. Use it when operating "
    "or demonstrating the stack."
)
H("First Boot", level=2)
CODE(
    "cd Week_02\\Day_10\\Capstone_Files\n"
    "Copy-Item .env.example .env\n"
    ".\\bootstrap.ps1\n"
    "python tests\\smoke_test.py\n"
)
H("Demo Flow", level=2)
OL([
    "Open the live map at http://localhost:5003 and show aircraft, weather, and seismic tabs.",
    "Open Kafka UI at http://localhost:8088 and show live stream topics plus config.public.* CDC topics.",
    "Open MinIO at http://localhost:9001 and show bronze Avro, silver Avro, and gold Parquet paths.",
    "Open Airflow at http://localhost:8080 and show 30_hop_medallion, 40_data_quality, and 50_business_kpis.",
    "Open the business dashboard at http://localhost:5002 and explain the four gold marts.",
    "Open the quality dashboard at http://localhost:5001 and explain rule results, DLQ sizes, and alerts.",
    "Open Grafana at http://localhost:3000 for pipeline metrics.",
])
H("Common Operations", level=2)
UL([
    "Tail an ingestor: docker logs -f ingestor-ogn (or ingestor-noaa / ingestor-seismic).",
    "Restart ingestors: docker compose restart ingestor-ogn ingestor-noaa ingestor-seismic.",
    "Force a gold mart: docker exec hop python3 /files/project/transforms/silver_to_gold.py --mart=seismic_24h_summary.",
    "Watch replication lag: docker exec config-db psql -U postgres -d config -c \"SELECT slot_name, active, restart_lsn FROM pg_replication_slots;\"",
    "Tear down without data loss: docker compose down.",
    "Clean slate: docker compose down -v.",
])
H("Troubleshooting", level=2)
UL([
    "bootstrap.ps1 waits forever for Kafka Connect: check docker logs connect and re-run debezium/download_plugins.ps1.",
    "noaa.observations is empty after 5+ minutes: set a real contact value in NOAA_USER_AGENT and recreate ingestor-noaa.",
    "seismic.events is quiet: verify the EMSC live page; quiet periods are possible.",
    "live map disconnected: refresh the page and check docker logs live-map-dashboard.",
    "Connect task FAILED: check plugin path, Schema Registry health, and connector status at http://localhost:8083/connectors.",
    "gold is empty: trigger 30_hop_medallion manually; low message volume can delay bronze flushes.",
    "Airflow webserver 502 on first hit: wait for airflow-init migration and check docker compose logs airflow-init.",
])

# ---- Submission ----------------------------------------------------------
H("Submission Checklist")
UL([
    "Architecture explanation from Appendix A — annotate or replace the diagram if your implementation differs.",
    "Screenshot pack: Kafka UI topics, MinIO bronze tree, MinIO gold mart, business dashboard, live map (with selected plane + magenta trail), quality dashboard during a failure.",
    "Output of python tests/smoke_test.py with all checks green.",
    "One sentence KPI definition + two dashboard artefacts (live map + business dashboard count as two).",
    "Short operations note for each of the three Task 7 failures: incident → evidence → resolution → prevention.",
    "When finished: docker compose down. Use docker compose down -v only if you intentionally want to erase MinIO + Postgres volumes.",
])

# ---- References ----------------------------------------------------------
H("References inside Capstone_Files/")
UL([
    "README.md — project overview and URLs.",
    "../Capstone_Guide/Capstone.docx — canonical student lab guide, architecture reference, and runbook.",
    "../Capstone_Guide/Capstone_Run_Guide.md — optional detailed operator walkthrough.",
    "bootstrap.ps1 / bootstrap.sh — idempotent stack bring-up.",
    "tests/smoke_test.py — end-to-end smoke test.",
])

doc.save(OUT)
print(f"wrote {OUT}")
