"""One-shot generator for Capstone_Guide/Capstone.docx.

Renders the Day-10 capstone architecture as a text/ASCII diagram inside
a Word document. No matplotlib — just python-docx + monospaced runs.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"C:\Users\Gamer\Documents\GitHub\Medallion pipeline\Week_02\Day_10\Capstone_Guide\Capstone.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# Page margins so the wide ASCII diagram doesn't wrap.
for section in doc.sections:
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

# ---- Title ------------------------------------------------------------------
title = doc.add_heading("Day 10 Capstone — Architecture", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Real streams  →  Kafka  →  Medallion on MinIO  →  Dashboards")
r.italic = True
r.font.size = Pt(11)

doc.add_paragraph()

# ---- Helper: write a fenced monospaced block --------------------------------
def mono_block(lines, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(size)

# =============================================================================
# Section 1: End-to-end flow
# =============================================================================
doc.add_heading("1. End-to-end data flow", level=1)
doc.add_paragraph(
    "Three independent public streams plus Postgres CDC feed Kafka. "
    "Kafka Connect sinks every topic to the bronze layer in MinIO. "
    "Apache Hop / Python transforms promote bronze → silver → gold. "
    "Two Flask dashboards read the gold marts and emit Prometheus metrics."
)

mono_block([
"+---------------------+    +---------------------+    +----------------------+",
"|  OGN APRS  (TCP)    |    |  NOAA api.weather   |    |  EMSC seismic  (WS)  |",
"|  ingestor-ogn       |    |  ingestor-noaa      |    |  ingestor-seismic    |",
"+----------+----------+    +----------+----------+    +-----------+----------+",
"           |                          |                           |           ",
"           |  ogn.aircraft.positions  |  noaa.observations        |  seismic.events",
"           |                          |  noaa.alerts              |           ",
"           v                          v                           v           ",
"  +===========================================================================+",
"  |                          KAFKA  (3 broker topics shown)                   |",
"  |                          + Schema Registry (Avro)                         |",
"  +======================================+====================================+",
"           ^                              |                                    ",
"           |  config.public.*  (CDC)      |  ALL topics                        ",
"           |                              v                                    ",
"  +--------+----------+        +----------+----------+                         ",
"  |  config-db (PG)   |        |  Kafka Connect      |                         ",
"  |  wal_level=logical|---WAL->|  - Debezium PG src  |                         ",
"  |  regions / thresh |        |  - S3 sink (CDC)    |                         ",
"  |  watchlist        |        |  - S3 sink (streams)|                         ",
"  +-------------------+        +----------+----------+                         ",
"                                          |                                    ",
"                                          v   Avro files, hourly partitions    ",
"                              +-----------+-----------+                        ",
"                              |   MinIO  s3://bronze  |                        ",
"                              +-----------+-----------+                        ",
"                                          |                                    ",
"                              hop/transforms/bronze_to_silver.py               ",
"                              (dedupe on PK, drop late >24h, audit)            ",
"                                          v                                    ",
"                              +-----------+-----------+                        ",
"                              |   MinIO  s3://silver  |  (Avro)                ",
"                              +-----------+-----------+                        ",
"                                          |                                    ",
"                              hop/transforms/silver_to_gold.py                 ",
"                              (DuckDB over silver Avro -> 4 marts)             ",
"                                          v                                    ",
"                              +-----------+-----------+                        ",
"                              |   MinIO  s3://gold    |  (Parquet)             ",
"                              +-----+-----------+-----+                        ",
"                                    |           |                              ",
"                                    v           v                              ",
"                    +---------------+--+     +--+----------------+             ",
"                    | quality-dashboard|     | business-dashboard|             ",
"                    |   :5001          |     |   :5002           |             ",
"                    +------------------+     +-------------------+             ",
])

doc.add_paragraph()

# =============================================================================
# Section 2: Orchestration (Airflow DAGs)
# =============================================================================
doc.add_heading("2. Orchestration — Airflow DAGs", level=1)
doc.add_paragraph(
    "The three stream ingestors are long-running services, not DAGs. "
    "Everything else is scheduled out of Airflow:"
)

mono_block([
"  00_bootstrap        (manual)    create topics, buckets, schemas, connectors",
"  15_config_drift     (every 2m)  mutate config tables -> keeps Debezium busy",
"  30_hop_medallion    (every 5m)  silver streams || silver CDC -> 4 gold marts",
"  40_data_quality     (every 5m)  rule pack -> POST quality-dashboard",
"  50_business_kpis    (every 5m)  POST business-dashboard /api/refresh",
])

doc.add_paragraph()

# =============================================================================
# Section 3: Medallion layout
# =============================================================================
doc.add_heading("3. Medallion contracts in MinIO", level=1)

mono_block([
"  s3://bronze/<topic>/year=YYYY/month=MM/day=DD/hour=HH/<topic>+<part>+<off>.avro",
"  s3://silver/<entity>/year=YYYY/month=MM/day=DD/part-<hash>.avro",
"  s3://silver/_quality/<entity>/<utc-iso>.avro",
"  s3://gold/<mart>/snapshot=YYYYMMDDTHHMMSSZ/part-0.parquet",
"  s3://gold/<mart>/latest.parquet",
])

doc.add_paragraph()
doc.add_paragraph("Gold marts:")
mono_block([
"  aircraft_density_by_region   ogn_positions  + regions",
"  weather_snapshot             noaa_obs       + regions (joined on station_code)",
"  seismic_24h_summary          seismic_events bucketed magnitude x region",
"  region_alert_correlation     seismic + regions + thresholds + watchlist",
])

doc.add_paragraph()

# =============================================================================
# Section 4: Observability
# =============================================================================
doc.add_heading("4. Observability", level=1)

mono_block([
"  Kafka JMX           :9404 --+",
"  postgres-exporter   :9187 --+",
"  MinIO /metrics      :9000 --+      rule_files: data_quality, pipeline",
"  quality-dashboard   :5001 --+--> Prometheus :9090 --> Alertmanager :9093 --+",
"  business-dashboard  :5002 --+                                              |",
"  kafka-connect REST  :8083 --+                                              v",
"                                              quality-dashboard /webhook/alerts",
])

doc.add_paragraph()

# =============================================================================
# Section 5: Container inventory
# =============================================================================
doc.add_heading("5. Container inventory", level=1)
rows = [
    ("zookeeper",          "Kafka coordination"),
    ("kafka",              "Broker + JMX exporter :9404"),
    ("schema-registry",    "Confluent Schema Registry (Avro)"),
    ("connect",            "Kafka Connect (Debezium PG + 2x S3 sink)"),
    ("kafka-ui",           "Provectus Kafka UI"),
    ("config-db",          "Postgres 15, wal_level=logical (reference tables)"),
    ("postgres-exporter",  "Prometheus exporter for config-db"),
    ("app",                "Idle utility container (docker exec target)"),
    ("ingestor-ogn",       "OGN APRS -> ogn.aircraft.positions"),
    ("ingestor-noaa",      "NOAA REST -> noaa.observations / noaa.alerts"),
    ("ingestor-seismic",   "EMSC websocket -> seismic.events"),
    ("minio",              "S3-compatible storage (bronze/silver/gold)"),
    ("hop",                "Apache Hop GUI + python sidecar"),
    ("airflow-db",         "Postgres 15 (Airflow metadata)"),
    ("airflow-init",       "airflow db migrate + admin user (run-once)"),
    ("airflow-webserver",  "Airflow UI :8080"),
    ("airflow-scheduler",  "DAG scheduler"),
    ("quality-dashboard",  "Flask DataOps view, /metrics, alert webhook"),
    ("business-dashboard", "Flask reading gold parquet, /metrics"),
    ("prometheus",         "Metrics + rule evaluation"),
    ("alertmanager",       "Routes alerts -> quality dashboard webhook"),
    ("grafana",            "Provisioned dashboards"),
]
table = doc.add_table(rows=1 + len(rows), cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Container"
hdr[1].text = "Purpose"
for i, (c, purpose) in enumerate(rows, start=1):
    table.rows[i].cells[0].text = c
    table.rows[i].cells[1].text = purpose

doc.add_paragraph()
doc.add_paragraph(
    "All services attach to the bridge network `day10_pipeline`.",
)

doc.save(OUT)
print(f"wrote: {OUT}")
